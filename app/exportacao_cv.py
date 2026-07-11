"""Exportação do currículo padronizado para Word (.docx).

Gera um `.docx` no **mesmo layout-padrão** que a tool `estruturar_cv` reconhece
(contato no topo, seções em maiúsculas, `•` por curso/skill/idioma/certificação e
`➢` por vaga com empresa/cargo/período em linhas separadas). Assim o arquivo
baixado é editável no Word **mantendo o padrão** e, se reenviado, volta a ser
extraído corretamente (round-trip). Ver docs/dicionario_dados_curriculo_estruturado.md.
"""
from __future__ import annotations

import io

from agents.modelos import CurriculoEstruturado


def curriculo_para_docx(cv: CurriculoEstruturado) -> bytes:
    """Serializa o CV padronizado como um documento Word (.docx) em memória."""
    import docx
    from docx.shared import Pt

    doc = docx.Document()

    def _titulo(texto: str) -> None:
        run = doc.add_paragraph().add_run(texto)
        run.bold = True
        run.font.size = Pt(16)

    def _secao(texto: str) -> None:
        run = doc.add_paragraph().add_run(texto)
        run.bold = True

    # Cabeçalho: nome + contato (localização, telefone, e-mail, LinkedIn).
    dp = cv.dados_pessoais
    _titulo(dp.nome.strip() or "Currículo padronizado")
    if dp.localizacao.strip():
        doc.add_paragraph(dp.localizacao.strip())
    if dp.telefone.strip():
        doc.add_paragraph(dp.telefone.strip())
    if dp.email.strip():
        doc.add_paragraph(f"Email: {dp.email.strip()}")
    if dp.linkedin.strip():
        doc.add_paragraph(f"LinkedIn: {dp.linkedin.strip()}")

    if cv.resumo.strip():
        _secao("RESUMO PROFISSIONAL")
        doc.add_paragraph(cv.resumo.strip())

    if cv.formacoes_validas():
        _secao("FORMAÇÃO ACADÊMICA")
        for f in cv.formacoes_validas():
            partes = ", ".join(p.strip() for p in (f.curso, f.instituicao) if p.strip())
            periodo = f" ({f.periodo.strip()})" if f.periodo.strip() else ""
            doc.add_paragraph(f"• {partes}{periodo}")

    certs = [c.strip() for c in cv.certificacoes if c.strip()]
    if certs:
        _secao("CERTIFICAÇÕES")
        for c in certs:
            doc.add_paragraph(f"• {c}")

    if cv.skills_validas():
        _secao("COMPETÊNCIAS TÉCNICAS (STACK)")
        for s in cv.skills_validas():
            doc.add_paragraph(f"• {s}")

    if cv.idiomas_validos():
        _secao("IDIOMA")
        for i in cv.idiomas_validos():
            doc.add_paragraph(f"• {i}")

    if cv.experiencias_validas():
        _secao("EXPERIÊNCIA PROFISSIONAL")
        for e in cv.experiencias_validas():
            doc.add_paragraph(f"➢ {e.empresa.strip()}")
            if e.cargo.strip():
                doc.add_paragraph(e.cargo.strip())
            if e.periodo.strip():
                doc.add_paragraph(e.periodo.strip())
            if e.descricao.strip():
                doc.add_paragraph(e.descricao.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
