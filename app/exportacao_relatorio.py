"""Exportação do relatório de match CV × vaga para Word (.docx).

Gera, em memória, um `.docx` com o mesmo conteúdo do dashboard da tela **Nova
análise**: scores (ATS/aprofundado), cobertura de must-haves, resumo/highlights,
requisitos obrigatórios e gaps por prioridade. Espelha o padrão de
`app/exportacao_cv.py` (import lazy de `docx`, helpers de título/seção,
`io.BytesIO` → `bytes`). **Sem recomendações** nesta etapa.
"""
from __future__ import annotations

import io

from agents.modelos import AnaliseCV, ProjetoRecomendado, SugestaoSecao

_ORDEM_PRIORIDADE = ["ALTA", "MÉDIA", "BAIXA"]


def relatorio_para_docx(
    analise: AnaliseCV,
    empresa: str = "",
    cargo: str = "",
    *,
    sugestoes_cv: list[SugestaoSecao] | None = None,
    projetos: list[ProjetoRecomendado] | None = None,
) -> bytes:
    """Serializa o relatório do match como um documento Word (.docx) em memória.

    `sugestoes_cv` e `projetos` são opcionais: quando fornecidos (tela Sugestões),
    acrescentam as seções de recomendação (cursos/certificações para fechar os
    gaps, reescrita do CV por seção e projetos STAR a citar).
    """
    import docx
    from docx.shared import Pt

    doc = docx.Document()

    def _titulo(texto: str) -> None:
        run = doc.add_paragraph().add_run(texto)
        run.bold = True
        run.font.size = Pt(16)

    def _secao(texto: str) -> None:
        run = doc.add_paragraph().add_run(texto)
        run.bold = True

    # Cabeçalho.
    _titulo("Relatório do match CV × vaga")
    alvo = " — ".join(p.strip() for p in (empresa, cargo) if p and p.strip())
    if alvo:
        doc.add_paragraph(alvo)

    # Scores.
    atendidos, total, pct = analise.cobertura_must_have()
    _secao("SCORES")
    doc.add_paragraph(f"• Score aprofundado: {analise.score_aprofundado}/100")
    doc.add_paragraph(f"• Score ATS: {analise.score_ats}/100")
    doc.add_paragraph(f"• Cobertura de must-haves: {atendidos}/{total} ({pct}%)")

    # Resumo + highlights.
    if analise.resumo.strip():
        _secao("RESUMO")
        doc.add_paragraph(analise.resumo.strip())
    highlights = [
        analise.highlight_aprofundado,
        analise.highlight_ats,
        analise.highlight_must_have,
    ]
    highlights = [h.strip() for h in highlights if h and h.strip()]
    if highlights:
        _secao("HIGHLIGHTS POR DIMENSÃO")
        for h in highlights:
            doc.add_paragraph(f"• {h}")

    # Requisitos obrigatórios (must-haves).
    if analise.must_haves:
        _secao("REQUISITOS OBRIGATÓRIOS")
        for m in analise.must_haves:
            marca = "[OK]" if m.atende else "[FALTA]"
            linha = f"{marca} {m.requisito}"
            if m.atende and m.evidencia.strip():
                linha += f" — {m.evidencia.strip()}"
            doc.add_paragraph(linha)

    # Gaps por prioridade.
    if analise.gaps:
        _secao("GAPS POR PRIORIDADE")
        for prioridade in _ORDEM_PRIORIDADE:
            do_nivel = [g for g in analise.gaps if g.prioridade == prioridade]
            for g in do_nivel:
                doc.add_paragraph(f"• ({g.prioridade}) {g.titulo} — {g.descricao}".rstrip(" —"))

    # Recomendações para fechar os gaps (cursos/certificações + PoC).
    gaps_com_rec = [
        g for g in analise.gaps
        if g.recomendacao.strip() or g.cursos_certificacoes or g.projetos_portfolio
    ]
    if gaps_com_rec:
        _secao("RECOMENDAÇÕES PARA FECHAR OS GAPS")
        for prioridade in _ORDEM_PRIORIDADE:
            for g in (g for g in gaps_com_rec if g.prioridade == prioridade):
                doc.add_paragraph(f"[{g.prioridade}] {g.titulo}")
                if g.recomendacao.strip():
                    doc.add_paragraph(f"    Recomendação: {g.recomendacao.strip()}")
                for c in g.cursos_certificacoes:
                    doc.add_paragraph(f"    • Curso/Certificação: {c}")
                for p in g.projetos_portfolio:
                    doc.add_paragraph(f"    • Projeto sugerido: {p}")

    # Reescrita do CV por seção (só quando gerada na tela Sugestões).
    if sugestoes_cv:
        _secao("REESCRITA DO CV POR SEÇÃO")
        for s in sugestoes_cv:
            doc.add_paragraph(f"[{s.secao}]")
            doc.add_paragraph(f"    Original: {s.original}")
            doc.add_paragraph(f"    Sugestão: {s.sugestao}")
            if s.justificativa.strip():
                doc.add_paragraph(f"    Ações aplicadas: {s.justificativa.strip()}")
            if s.palavras_chave.strip():
                doc.add_paragraph(f"    Palavras-chave ATS: {s.palavras_chave.strip()}")

    # Projetos STAR a citar (recomendados a partir do portfólio cadastrado).
    if projetos:
        _secao("PROJETOS STAR A CITAR")
        for r in projetos:
            doc.add_paragraph(f"• {r.projeto} — {r.motivo}".rstrip(" —"))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
